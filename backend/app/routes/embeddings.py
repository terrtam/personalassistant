from io import BytesIO
from pathlib import Path
from typing import Any

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from pydantic import BaseModel, Field

from app.core.settings import get_settings
from app.services.embeddings.pipeline import build_index, search_index

router = APIRouter(prefix="/embeddings", tags=["embeddings"])


class EmbeddingDocument(BaseModel):
    id: str | None = None
    text: str = Field(..., min_length=1)
    metadata: dict[str, Any] = Field(default_factory=dict)


class BuildIndexRequest(BaseModel):
    documents: list[EmbeddingDocument] = Field(..., min_length=1)
    chunk_size: int | None = Field(default=None, ge=100, le=4000)
    chunk_overlap: int | None = Field(default=None, ge=0, le=1000)


class BuildIndexResponse(BaseModel):
    provider: str
    model: str
    index_path: str
    source_documents: int
    chunks: int


class SearchRequest(BaseModel):
    query: str = Field(..., min_length=1)
    k: int = Field(default=5, ge=1, le=20)


class SearchResult(BaseModel):
    text: str
    metadata: dict[str, Any]
    score: float


class SearchResponse(BaseModel):
    results: list[SearchResult]


def _extract_pdf_text(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="PDF upload requires 'pypdf'. Install it in backend/requirements.txt.",
        ) from exc

    reader = PdfReader(BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts).strip()


def _extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="DOCX upload requires 'python-docx'. Install it in backend/requirements.txt.",
        ) from exc

    document = Document(BytesIO(content))
    return "\n".join(p.text for p in document.paragraphs).strip()


def _extract_upload_text(filename: str, content: bytes) -> tuple[str, str]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(content), "pdf"
    if suffix == ".docx":
        return _extract_docx_text(content), "docx"
    if suffix == ".txt":
        return content.decode("utf-8", errors="ignore").strip(), "txt"
    raise HTTPException(
        status_code=400,
        detail="Unsupported file type. Allowed: .pdf, .docx, .txt",
    )


@router.post("/index", response_model=BuildIndexResponse)
async def index_embeddings(payload: BuildIndexRequest) -> BuildIndexResponse:
    try:
        counts = build_index(
            documents=[doc.model_dump() for doc in payload.documents],
            chunk_size=payload.chunk_size,
            chunk_overlap=payload.chunk_overlap,
        )
        settings = get_settings()
        return BuildIndexResponse(
            provider=settings.embedding_provider,
            model=settings.embedding_model,
            index_path=settings.embedding_index_path,
            source_documents=counts["source_documents"],
            chunks=counts["chunks"],
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to build embedding index: {str(exc)}",
        ) from exc


@router.post("/search", response_model=SearchResponse)
async def search_embeddings(payload: SearchRequest) -> SearchResponse:
    try:
        results = search_index(query=payload.query, k=payload.k)
        return SearchResponse(
            results=[SearchResult(**item) for item in results],
        )
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to search embedding index: {str(exc)}",
        ) from exc


@router.post("/upload", response_model=BuildIndexResponse)
async def upload_embeddings(
    file: UploadFile = File(...),
    chunk_size: int | None = Form(default=None),
    chunk_overlap: int | None = Form(default=None),
) -> BuildIndexResponse:
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="Filename is required.")

        content = await file.read()
        text, extracted_from = _extract_upload_text(file.filename, content)
        if not text:
            raise HTTPException(
                status_code=400,
                detail="Uploaded file has no extractable text content.",
            )

        counts = build_index(
            documents=[
                {
                    "id": file.filename,
                    "text": text,
                    "metadata": {
                        "filename": file.filename,
                        "content_type": file.content_type,
                        "source_type": extracted_from,
                    },
                }
            ],
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        settings = get_settings()
        return BuildIndexResponse(
            provider=settings.embedding_provider,
            model=settings.embedding_model,
            index_path=settings.embedding_index_path,
            source_documents=counts["source_documents"],
            chunks=counts["chunks"],
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to upload/index file: {str(exc)}",
        ) from exc
